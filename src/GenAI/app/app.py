from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import List, Optional, Literal
from pathlib import Path
import json
import base64

from GenAI.clients.openai_client import openai_client
from GenAI.clients.gemini_client import gemini_client


APP_DIR = Path(__file__).parent
try:
    # repo root if layout is repo_root/src/GenAI/app/app.py
    ROOT_DIR = Path(__file__).resolve().parents[3]
except IndexError:
    ROOT_DIR = APP_DIR
OUT_DIR = (ROOT_DIR / "data" / "outputs")
OUT_DIR.mkdir(parents=True, exist_ok=True)


class BriefParams(BaseModel):
    project_name: str = "Proyecto Sin Nombre"
    objective_primary: str = "Incrementar awareness"
    audience_description: str = "B2B decisores LATAM"
    segments: List[str] = ["CMOs", "Marketing Managers"]
    channels: List[str] = ["LinkedIn", "Email"]
    deliverables: List[str] = ["KV", "Landing", "3 piezas social"]
    tone: List[str] = ["Profesional", "Claro"]
    budget_total: str = "USD 10,000"
    deadline_start: str = "2025-11-01"
    deadline_end: str = "2025-12-01"
    extra: Optional[dict] = None
    language: Literal["es", "en"] = "es"
    export_docx: bool = False


class TTSParams(BaseModel):
    text: str = "Hola Julio, todo ok."
    voice: Literal["verse", "alloy", "soft", "sage"] = "verse"
    filename: str = "voz.mp3"


class ImageParams(BaseModel):
    prompt: str
    size: Literal["512x512", "1024x1024", "2048x2048"] = "1024x1024"
    provider: Literal["openai", "gemini"] = "openai"
    filename: str = "image.png"


app = FastAPI(title="Globe Creative Engine (MVP)")


BRIEF_SCHEMA = {
    "project_name": "string",
    "background": "string",
    "objective_primary": "string",
    "objectives_secondary": ["string"],
    "target_audience": {"description": "string", "segments": ["string"]},
    "market_context": "string",
    "product_value": "string",
    "unique_selling_proposition": "string",
    "key_message": "string",
    "tone_style": ["string"],
    "deliverables": ["string"],
    "channels": ["string"],
    "mandatories": {"brand_assets": ["string"], "legal": ["string"]},
    "creative_directions": ["string"],
    "timeline": {
        "start_date": "YYYY-MM-DD",
        "end_date": "YYYY-MM-DD",
        "milestones": [{"name": "string", "date": "YYYY-MM-DD"}],
    },
    "budget": {
        "total": "number|string",
        "breakdown": [{"item": "string", "amount": "number|string"}],
    },
    "kpis_primary": ["string"],
    "kpis_secondary": ["string"],
    "cta": "string",
    "risks_assumptions": ["string"],
    "approvals_stakeholders": ["string"],
    "references_links": ["string"],
    "notes": "string",
}


def schema_str(d, indent=0):
    sp = "  " * indent
    if isinstance(d, dict):
        inner = ",\n".join([f'{sp}  "{k}": {schema_str(v, indent+1)}' for k, v in d.items()])
        return "{\n" + inner + f"\n{sp}" + "}"
    if isinstance(d, list):
        return "[" + schema_str(d[0], indent + 1) + "]" if d else '["string"]'
    return f'"{d}"'


def build_brief_prompt(params: BriefParams):
    sys = (
        "Eres estratega creativo senior. Devuelve SOLO JSON válido conforme al esquema."
    )
    instr = (
        f"Esquema:\n{schema_str(BRIEF_SCHEMA)}\n"
        f"Rellena todos los campos de forma concisa y accionable. "
        f"Si falta info, infiérela y marca en 'notes'. Idioma: {params.language}."
    )
    payload = {
        "project_name": params.project_name,
        "objective_primary": params.objective_primary,
        "target_audience": {
            "description": params.audience_description,
            "segments": params.segments,
        },
        "channels": params.channels,
        "deliverables": params.deliverables,
        "tone_style": params.tone,
        "timeline": {
            "start_date": params.deadline_start,
            "end_date": params.deadline_end,
        },
        "budget": {"total": params.budget_total},
    }
    if params.extra:
        payload.update(params.extra)
    return [
        {"role": "system", "content": sys},
        {"role": "user", "content": instr},
        {"role": "user", "content": json.dumps({"params": payload}, ensure_ascii=False)},
    ]


def save_docx(data: dict, out_path: Path):
    try:
        from docx import Document
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"python-docx no instalado: {e}")
    doc = Document()
    doc.add_heading(data.get("project_name", "Brief"), 0)

    def add(h, b):
        doc.add_heading(h, 1)
        if isinstance(b, list):
            for x in b:
                doc.add_paragraph(f"• {x}")
        elif isinstance(b, dict):
            for k, v in b.items():
                doc.add_paragraph(f"{k}: {v if not isinstance(v, (list, dict)) else ''}")
                if isinstance(v, list):
                    for x in v:
                        doc.add_paragraph(f"   - {x}")
                if isinstance(v, dict):
                    for kk, vv in v.items():
                        doc.add_paragraph(f"   - {kk}: {vv}")
        else:
            doc.add_paragraph(str(b))

    sections = [
        ("Antecedentes", data.get("background")),
        ("Objetivo principal", data.get("objective_primary")),
        ("Objetivos secundarios", data.get("objectives_secondary", [])),
        ("Audiencia", data.get("target_audience", {})),
        ("Contexto", data.get("market_context")),
        ("Propuesta de valor", data.get("product_value")),
        ("USP", data.get("unique_selling_proposition")),
        ("Mensaje clave", data.get("key_message")),
        ("Tono/Estilo", data.get("tone_style", [])),
        ("Entregables", data.get("deliverables", [])),
        ("Canales", data.get("channels", [])),
        ("Mandatorios", data.get("mandatories", {})),
        ("Direcciones creativas", data.get("creative_directions", [])),
        ("Timeline", data.get("timeline", {})),
        ("Budget", data.get("budget", {})),
        ("KPIs primarios", data.get("kpis_primary", [])),
        ("KPIs secundarios", data.get("kpis_secondary", [])),
        ("CTA", data.get("cta")),
        ("Riesgos & supuestos", data.get("risks_assumptions", [])),
        ("Stakeholders", data.get("approvals_stakeholders", [])),
        ("Referencias", data.get("references_links", [])),
        ("Notas", data.get("notes")),
    ]
    for h, b in sections:
        if b:
            add(h, b)
    doc.save(out_path)


@app.post("/brief")
def create_brief(p: BriefParams):
    gpt = openai_client()
    msgs = build_brief_prompt(p)
    resp = gpt.chat.completions.create(
        model="gpt-4o-mini", messages=msgs, response_format={"type": "json_object"}, temperature=0.2
    )
    data = json.loads(resp.choices[0].message.content)

    result = {"brief": data}
    if p.export_docx:
        out = OUT_DIR / f"brief_{p.project_name.replace(' ', '_')}.docx"
        save_docx(data, out)
        result["docx_path"] = str(out)
    return result


@app.post("/tts")
def tts(p: TTSParams):
    gpt = openai_client()
    out = OUT_DIR / p.filename
    with gpt.audio.speech.with_streaming_response.create(
        model="gpt-4o-mini-tts", voice=p.voice, input=p.text
    ) as r:
        r.stream_to_file(out)
    return {"file": str(out)}


@app.post("/image")
def image(p: ImageParams):
    out = OUT_DIR / p.filename
    if p.provider == "openai":
        gpt = openai_client()
        r = gpt.images.generate(model="gpt-image-1", prompt=p.prompt, size=p.size)
        b64 = r.data[0].b64_json
        out.write_bytes(base64.b64decode(b64))
        return {"file": str(out), "provider": "openai"}
    else:
        gem = gemini_client()
        r = gem.models.generate_content(model="gemini-2.5-flash-image", contents=[p.prompt])
        # buscar binario
        for c in r.candidates:
            for part in c.content.parts:
                if getattr(part, "inline_data", None):
                    from io import BytesIO
                    from PIL import Image

                    im = Image.open(BytesIO(part.inline_data.data))
                    im.save(out)
                    return {"file": str(out), "provider": "gemini"}
        raise HTTPException(status_code=500, detail="Gemini no retornó imagen")


@app.get("/download")
def download(path: str):
    # Restringe descargas al directorio OUT_DIR para evitar path traversal
    base = OUT_DIR.resolve()
    fp = (OUT_DIR / path).resolve()
    try:
        if not fp.is_relative_to(base):
            raise HTTPException(403, "Solo se permite descargar desde outputs")
    except AttributeError:
        if str(fp)[: len(str(base))] != str(base):
            raise HTTPException(403, "Solo se permite descargar desde outputs")
    if not fp.exists():
        raise HTTPException(404, "Archivo no encontrado")
    return FileResponse(fp)

