from pathlib import Path
from docx import Document

from GenAI.clients.openai_client import openai_client


def main():
    client = openai_client()
    prompt = (
        "Crea un mini informe titulado 'Tendencias de Inteligencia Artificial 2025', "
        "con secciones: Introduccion, Aplicaciones Clave, Desafios y Conclusiones. "
        "Escribe en tono profesional y conciso."
    )
    resp = client.chat.completions.create(
        model="gpt-4o", messages=[{"role": "user", "content": prompt}]
    )
    texto = resp.choices[0].message.content
    doc = Document()
    doc.add_heading("Tendencias de Inteligencia Artificial 2025", 0)
    doc.add_paragraph(texto)
    out = Path("Tendencias_IA_2025.docx").resolve()
    doc.save(out)
    print(f"Documento generado: {out}")


if __name__ == "__main__":
    main()

